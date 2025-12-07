import uuid
from typing import List
from fastapi import UploadFile
from app.services.vector_store import get_vector_store
from app.services.llm_service import llm_service
from app.models.document import Document
from sqlalchemy.ext.asyncio import AsyncSession
from pypdf import PdfReader
import docx
import io

class IngestionService:
    def __init__(self, db: AsyncSession):
        self.db = db
        self.vector_store = get_vector_store()

    async def process_file(self, file: UploadFile) -> Document:
        content = await file.read()
        file_size = len(content)
        
        # 1. Extract Text
        text = await self._extract_text(content, file.filename, file.content_type)
        
        # 2. Chunk Text
        chunks = self._chunk_text(text)
        
        # 3. Create Document Record
        doc = Document(
            filename=file.filename,
            content_type=file.content_type,
            file_size=file_size,
            chunk_count=len(chunks)
        )
        self.db.add(doc)
        await self.db.flush()
        
        # 4. Embed and Store
        embeddings = []
        ids = []
        metadatas = []
        documents_for_vector_db = []
        
        # Track current section statefully across chunks
        current_section = "Introduction"  # Default start
        
        for i, chunk in enumerate(chunks):
            # Check if this chunk starts a new section
            possible_header = self._extract_section_header(chunk)
            if possible_header != "unknown":
                current_section = possible_header
            
            embedding = await llm_service.get_embedding(chunk)
            embeddings.append(embedding)
            
            chunk_id = f"{doc.id}_{i}"
            ids.append(chunk_id)
            
            metadatas.append({
                "document_id": str(doc.id),
                "chunk_index": i,
                "source": file.filename,
                "section": current_section,  # Use stateful section
                "chunk_length": len(chunk)
            })
            documents_for_vector_db.append(chunk)

        if documents_for_vector_db:
            self.vector_store.add_documents(
                documents=documents_for_vector_db,
                metadatas=metadatas,
                ids=ids,
                embeddings=embeddings
            )

        return doc

    def _extract_section_header(self, text: str) -> str:
        """
        Extract section header from chunk text.
        """
        lines = text.split('\n')
        for line in lines[:2]:  # Check first 2 lines
            line = line.strip()
            # Match numbered sections (1. Introduction) or All Caps Headers (BACKGROUND)
            if len(line) < 100 and (
                (line[0].isdigit() and "." in line[:5]) or 
                (line.isupper() and len(line) > 3) or 
                line.startswith("#")
            ):
                return line.lstrip('#').strip()
        return "unknown"

    async def _extract_text(self, content: bytes, filename: str, content_type: str) -> str:
        text = ""
        if content_type == "application/pdf" or filename.endswith(".pdf"):
            reader = PdfReader(io.BytesIO(content))
            for page in reader.pages:
                text += page.extract_text() + "\n"
        
        elif content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document" or filename.endswith(".docx"):
            doc = docx.Document(io.BytesIO(content))
            text = "\n".join([para.text for para in doc.paragraphs])
            
        elif content_type == "text/plain" or filename.endswith(".txt"):
            text = content.decode("utf-8")
        
        else:
            raise ValueError(f"Unsupported file type: {content_type}")
        
        return self._clean_text(text)

    def _clean_text(self, text: str) -> str:
        """
        Clean and normalize text using standard NLP practices.
        """
        import re
        # 1. Normalize line breaks (hyphenation at end of line)
        # e.g. "exam-\nple" -> "example"
        text = re.sub(r'-\n', '', text)
        
        # 2. Collapse excessive whitespace
        # "  " -> " "
        text = re.sub(r'\s+', ' ', text)
        
        # 3. Clean headers (e.g. "1 . Introduction" -> "1. Introduction")
        text = re.sub(r'(\d+)\s+\.', r'\1.', text)
        
        return text.strip()

    def _chunk_text(self, text: str, target_size: int = 800, overlap: int = 200) -> List[str]:
        """
        Recursive Character Text Splitter (Standard RAG Pattern)
        Priority splits: Paragraphs -> Lines -> Sentences -> Words
        """
        if not text:
            return []
            
        separators = ["\n\n", "\n", ". ", " ", ""]
        chunks = []
        
        def split_text(text: str, separators: List[str]) -> List[str]:
            """Recursively split text using the first valid separator."""
            final_chunks = []
            separator = separators[0]
            new_separators = separators[1:]
            
            # Split current text by separator
            splits = text.split(separator) if separator else list(text)
            
            good_splits = []
            
            # Merge small splits back together
            current_split = ""
            for s in splits:
                # Add separator back if it's not empty string
                to_add = s + separator if separator else s
                
                if len(current_split) + len(to_add) <= target_size:
                    current_split += to_add
                else:
                    if current_split:
                        good_splits.append(current_split)
                    current_split = to_add
            
            if current_split:
                good_splits.append(current_split)
                
            # If standard splitting failed (chunks still too big), recurse
            return good_splits

        # 1. Initial split by double newline (Paragraphs)
        paragraphs = text.replace('\r', '\n').split('\n\n')
        
        current_chunk = ""
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
            
            # If paragraph fits, add it
            if len(current_chunk) + len(para) + 2 <= target_size:
                current_chunk += "\n\n" + para if current_chunk else para
            else:
                # Chunk full, save it
                if current_chunk:
                    chunks.append(current_chunk)
                
                # If single paragraph is too big, use recursive splitting
                if len(para) > target_size:
                    # Paragraph -> Sentences -> Words
                    sub_chunks = split_text(para, [". ", " "])
                    
                    # Add all fully formed sub-chunks
                    chunks.extend(sub_chunks[:-1])
                    # Keep last part as start of next chunk
                    current_chunk = sub_chunks[-1]
                else:
                    current_chunk = para
        
        if current_chunk:
            chunks.append(current_chunk)
        
        # Add overlap
        if overlap > 0 and len(chunks) > 1:
            overlapped = []
            for i, chunk in enumerate(chunks):
                if i > 0:
                    prev = chunks[i-1]
                    # Take last 'overlap' chars, try to respect word boundary
                    overlap_text = prev[-overlap:]
                    first_space = overlap_text.find(" ")
                    if first_space != -1:
                        overlap_text = overlap_text[first_space+1:]
                    chunk = overlap_text + " " + chunk
                overlapped.append(chunk)
            return overlapped
        
        return chunks
        
        # Add the last chunk
        if current_chunk:
            chunks.append(current_chunk)
        
        # Add overlap for context continuity
        if overlap > 0 and len(chunks) > 1:
            overlapped = []
            for i, chunk in enumerate(chunks):
                if i > 0:
                    # Add last N characters from previous chunk
                    prev_chunk = chunks[i-1]
                    overlap_text = prev_chunk[-overlap:] if len(prev_chunk) > overlap else prev_chunk
                    # Try to find sentence boundary
                    last_period = overlap_text.rfind('. ')
                    if last_period > overlap // 2:
                        overlap_text = overlap_text[last_period + 2:]
                    chunk = overlap_text + " " + chunk
                overlapped.append(chunk)
            return overlapped
        
        return chunks
